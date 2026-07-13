"""
Extract plain text from uploaded files (PDF, DOCX, text formats) with OCR fallback for scans.
"""

from __future__ import annotations

import io
import logging
from dataclasses import dataclass
from typing import Optional

from fastapi import HTTPException, status

from app.core.config import settings

logger = logging.getLogger(__name__)


@dataclass
class ExtractionResult:
    text: str
    page_count: Optional[int] = None
    method: str = "unknown"


def _extract_pdf_pypdf(raw: bytes) -> tuple[str, int]:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(raw))
    page_count = len(reader.pages)
    parts = []
    for page in reader.pages:
        parts.append(page.extract_text() or "")
    return "\n".join(parts).strip(), page_count


def _extract_pdf_pymupdf(raw: bytes) -> tuple[str, int]:
    import fitz

    doc = fitz.open(stream=raw, filetype="pdf")
    page_count = doc.page_count
    parts = []
    for page in doc:
        parts.append(page.get_text("text") or "")
    doc.close()
    return "\n".join(parts).strip(), page_count


def _extract_pdf_ocr(raw: bytes, max_pages: int) -> tuple[str, int]:
    import fitz
    import pytesseract
    from PIL import Image

    doc = fitz.open(stream=raw, filetype="pdf")
    page_count = doc.page_count
    limit = min(page_count, max_pages)
    parts: list[str] = []
    for i in range(limit):
        page = doc[i]
        pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
        img = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
        parts.append(pytesseract.image_to_string(img) or "")
    doc.close()
    text = "\n".join(parts).strip()
    if page_count > limit:
        text += f"\n\n[OCR limited to first {limit} of {page_count} pages]"
    return text, page_count


def _ocr_available() -> bool:
    if not settings.document_ocr_enabled:
        return False
    try:
        import pytesseract

        pytesseract.get_tesseract_version()
        return True
    except Exception:
        return False


def _extract_pdf(raw: bytes) -> ExtractionResult:
    text, page_count = _extract_pdf_pypdf(raw)
    method = "pypdf"

    min_chars = max(32, page_count * 40)
    if len(text) < min_chars:
        try:
            alt, page_count = _extract_pdf_pymupdf(raw)
            if len(alt) > len(text):
                text, method = alt, "pymupdf"
        except Exception as exc:
            logger.debug("pymupdf extraction skipped: %s", exc)

    if len(text) < min_chars and _ocr_available():
        try:
            ocr_text, page_count = _extract_pdf_ocr(raw, settings.document_ocr_max_pages)
            if len(ocr_text) > len(text):
                text, method = ocr_text, "ocr"
        except Exception as exc:
            logger.warning("PDF OCR failed: %s", exc)

    return ExtractionResult(text=text, page_count=page_count, method=method)


def _extract_docx(raw: bytes) -> ExtractionResult:
    from docx import Document as DocxDocument

    doc = DocxDocument(io.BytesIO(raw))
    parts = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return ExtractionResult(text="\n".join(parts).strip(), page_count=None, method="docx")


def extract_text_from_upload(filename: str, raw: bytes) -> ExtractionResult:
    """Return extracted text and metadata for an uploaded file."""
    if not raw:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Empty file")

    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    if ext == "pdf":
        return _extract_pdf(raw)

    if ext in ("txt", "md", "csv", "html", "htm"):
        text = raw.decode("utf-8", errors="ignore").strip()
        return ExtractionResult(text=text, page_count=None, method=ext)

    if ext == "docx":
        return _extract_docx(raw)

    if ext == "doc":
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Legacy .doc files are not supported. Save as .docx or PDF and re-upload.",
        )

    raise HTTPException(
        status_code=status.HTTP_400_BAD_REQUEST,
        detail=f"Unsupported file type: .{ext or 'unknown'}",
    )
